"""
Generate a professional PRD (.docx) for the AI-powered Recruiting Platform,
divided into 3 development phases.

Source: Updated_AI_Recruiting_Platform_Feature_Spec - Copy.docx
Output: AIRIS_PRD_v1.0.docx (written next to this script's project root)

This generator intentionally keeps logic explicit and readable so that the
document structure, phasing, and content can be reviewed and edited in one
place before regenerating.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

BRAND_PRIMARY = RGBColor(0x0B, 0x3D, 0x91)   # deep blue
BRAND_ACCENT = RGBColor(0x1F, 0x6F, 0xEB)    # accent blue
MUTED = RGBColor(0x55, 0x5F, 0x6D)
HEADER_FILL = "0B3D91"
ROW_FILL = "EEF3FB"


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None,
                  size: int = 10, align: int = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_PRIMARY if level <= 1 else BRAND_ACCENT


def add_paragraph(doc: Document, text: str, *, bold: bool = False,
                  italic: bool = False, size: int = 11,
                  color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.runs[0] if p.runs else p.add_run("")
        run.text = item
        run.font.size = Pt(11)


def add_feature_table(doc: Document, headers: list[str], rows: list[list[str]],
                      col_widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, HEADER_FILL)
        set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                      size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Body rows
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            if r_idx % 2 == 0:
                shade_cell(cell, ROW_FILL)
            set_cell_text(cell, val, size=10)

    if col_widths_cm is not None:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing


def add_divider(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F6FEB")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------

def build_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI-Powered Recruiting Platform")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND_PRIMARY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Product Requirements Document (PRD)")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND_ACCENT

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Target users: Recruiting & staffing agencies (US and India)\n"
        "Version 1.0  |  April 2026  |  Status: Draft for review"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    doc.add_paragraph()
    add_divider(doc)
    doc.add_paragraph()


def build_executive_summary(doc: Document) -> None:
    add_heading(doc, "1. Executive summary", level=1)
    add_paragraph(
        doc,
        "This document translates the AI-Powered Recruiting Platform feature "
        "specification into a prioritised product requirements document (PRD). "
        "It defines the problem, target users, success metrics, functional and "
        "non-functional requirements, and a proposed three-phase delivery plan "
        "covering the first 18 months of development.",
    )
    add_paragraph(
        doc,
        "The platform is positioned as an AI-first SaaS for recruiting agencies "
        "operating across US and India markets. Its competitive wedge is the "
        "combination of agency-specific workflows (multi-client, white-label, "
        "placement-centric) with embedded AI for interviewing, scoring, and "
        "compliance \u2014 capabilities that generic ATS vendors and point-solution "
        "AI tools do not deliver in a single product.",
    )


def build_problem_and_users(doc: Document) -> None:
    add_heading(doc, "2. Problem statement & target users", level=1)

    add_heading(doc, "2.1 Problem", level=2)
    add_bullets(doc, [
        "Recruiting agencies juggle 10\u2013100+ concurrent clients, each with different "
        "processes, branding, and reporting expectations. Existing ATS tools were "
        "built for corporate HR, not agencies.",
        "Recruiters spend 40%+ of their time on candidate communication and "
        "scheduling \u2014 work that is highly automatable but fragmented across "
        "email, SMS, WhatsApp, and calendars.",
        "AI hiring tools (HireVue, Paradox, Mercor, etc.) solve narrow slices of "
        "the workflow; agencies must stitch 4\u20136 tools together, creating data "
        "silos and compliance gaps.",
        "Regulatory exposure is rising (NYC LL144, EU AI Act, India DPDPA). "
        "Agencies need auditable, bias-aware AI by default, not as an add-on.",
    ])

    add_heading(doc, "2.2 Primary personas", level=2)
    add_feature_table(
        doc,
        headers=["Persona", "Description", "Primary jobs-to-be-done"],
        rows=[
            ["Agency admin / owner",
             "Runs the agency; owns P&L, client relationships, and compliance",
             "Configure workspaces, manage users, monitor margins and placements"],
            ["Recruiter / account manager",
             "Sources, screens, and places candidates against client roles",
             "Intake roles, screen candidates, schedule interviews, submit shortlists"],
            ["Client hiring manager",
             "External user from the agency's client company",
             "Review shortlists, give feedback, track pipeline progress"],
            ["Candidate",
             "Job seeker engaging with the agency",
             "Apply, complete screening, schedule interviews, receive updates"],
        ],
        col_widths_cm=[3.5, 5.5, 7.0],
    )

    add_heading(doc, "2.3 Target markets", level=2)
    add_bullets(doc, [
        "United States: mid-market and enterprise staffing agencies; strong "
        "requirement for SOC 2, EEOC compliance, and NYC LL144 bias auditing.",
        "India: IT staffing and contract-hiring agencies; strong requirement for "
        "WhatsApp-first communication, DPDPA compliance, and integration with "
        "Keka / Darwinbox / Zoho Recruit.",
    ])


def build_goals_and_metrics(doc: Document) -> None:
    add_heading(doc, "3. Goals & success metrics", level=1)

    add_heading(doc, "3.1 Product goals", level=2)
    add_bullets(doc, [
        "Become the default operating system for AI-first recruiting agencies in "
        "the US and India within 18 months.",
        "Reduce recruiter time-per-hire by at least 40% through AI screening, "
        "scheduling automation, and unified communication.",
        "Deliver explainable, bias-audited AI that is defensible to enterprise "
        "procurement and regulators by design.",
    ])

    add_heading(doc, "3.2 Success metrics (North Star and supporting)", level=2)
    add_feature_table(
        doc,
        headers=["Metric", "Type", "Target (12 months post-launch)"],
        rows=[
            ["Placements processed per recruiter per month", "North Star", "+50% vs. baseline"],
            ["Time-to-shortlist (role opened \u2192 3 qualified candidates)",
             "Activation", "< 48 hours for 80% of roles"],
            ["AI screening coverage", "Engagement", "> 70% of inbound applications"],
            ["Interview no-show rate", "Efficiency", "< 10%"],
            ["Weekly active recruiters / seats sold", "Adoption", "> 75%"],
            ["Gross revenue retention", "Commercial", "> 95%"],
            ["SOC 2 Type II certification", "Compliance", "Achieved by end of Phase 3"],
            ["Adverse-impact audit pass rate", "Responsible AI", "100% of enterprise tenants"],
        ],
        col_widths_cm=[6.0, 4.0, 6.0],
    )


def build_scope(doc: Document) -> None:
    add_heading(doc, "4. Scope", level=1)
    add_paragraph(
        doc,
        "Scope is organised by the seven functional pillars from the source "
        "feature specification, plus the two AI tiers (essential and high-value). "
        "Each feature is tagged with its delivery phase in section 7. "
        "Anything not listed here is explicitly out of scope for the first 18 months.",
    )

    add_heading(doc, "4.1 Multi-client workspace management", level=2)
    add_feature_table(
        doc,
        headers=["Feature", "Description", "Why it matters"],
        rows=[
            ["Client workspaces",
             "Separate workspaces per client with custom branding, job templates, and access controls",
             "Data isolation between competing clients in the same industry"],
            ["White-label portals",
             "Client-facing portals with the agency's or client's branding, logo, and custom domain",
             "Agencies sell a branded service; generic platform branding undermines positioning"],
            ["Role-based access",
             "Granular permissions: agency admin, recruiter, client hiring manager, candidate",
             "Clients often want pipeline visibility without full platform access"],
            ["Multi-entity billing",
             "Track usage, interviews, and placements per client for invoicing and margin analysis",
             "Agencies bill per placement, retainer, or hourly"],
        ],
        col_widths_cm=[3.8, 6.2, 6.0],
    )

    add_heading(doc, "4.2 Candidate pipeline & relationship management", level=2)
    add_feature_table(
        doc,
        headers=["Feature", "Description", "Why it matters"],
        rows=[
            ["Universal candidate database",
             "Single searchable database across all clients with tagging, notes, and interaction history",
             "Agencies re-place candidates; losing context between engagements is costly"],
            ["Customisable pipelines",
             "Drag-and-drop Kanban boards with configurable stages per client or role",
             "Each client has different processes: some need 2 rounds, others 5"],
            ["Rich candidate profiles",
             "CV parsing, skills taxonomy, interview history, assessment scores, placement history",
             "Recruiters need a 360-degree view before submitting to clients"],
            ["Duplicate detection",
             "Automatic flagging of duplicate candidates across the database",
             "Duplicates waste time and embarrass the agency"],
            ["Communication hub",
             "Centralised email, SMS, and WhatsApp messaging with templates and tracking",
             "Recruiters spend 40%+ of time on candidate communication"],
            ["Submission tracking",
             "Track which candidates have been submitted to which clients, with status and feedback",
             "Prevents double-submissions and provides clear audit trails"],
        ],
        col_widths_cm=[3.8, 6.2, 6.0],
    )

    add_heading(doc, "4.3 Job management & distribution", level=2)
    add_feature_table(
        doc,
        headers=["Feature", "Description", "Why it matters"],
        rows=[
            ["Job requisition builder",
             "Structured intake forms for role requirements, skills, compensation, client preferences",
             "Standardises job intake; reduces back-and-forth with clients"],
            ["Application parsing",
             "Automatic CV/resume parsing into structured profiles with skills extraction",
             "Manual data entry wastes recruiter time and introduces errors"],
        ],
        col_widths_cm=[3.8, 6.2, 6.0],
    )

    add_heading(doc, "4.4 Interview scheduling & coordination", level=2)
    add_feature_table(
        doc,
        headers=["Feature", "Description", "Why it matters"],
        rows=[
            ["Automated scheduling",
             "Calendar sync (Google, Outlook, Exchange), timezone detection, self-service booking",
             "Recruiters lose 30 min\u20132 hr scheduling each interview"],
            ["Panel coordination",
             "Multi-interviewer availability matching, room booking, agenda distribution",
             "Agency-arranged panels with client teams are logistically complex"],
            ["Reminders & rescheduling",
             "Automated reminders via email, SMS, and WhatsApp with one-click reschedule",
             "Reduces no-shows by 30\u201350%"],
            ["Feedback collection",
             "Structured scorecard submission with deadline enforcement and nudges",
             "Timely feedback keeps pipelines moving"],
        ],
        col_widths_cm=[3.8, 6.2, 6.0],
    )

    add_heading(doc, "4.5 Assessment & evaluation infrastructure", level=2)
    add_feature_table(
        doc,
        headers=["Feature", "Description", "Why it matters"],
        rows=[
            ["Live video interviewing",
             "Built-in video with recording, transcription, and scorecard overlay",
             "Third-party tools fragment workflow"],
            ["Skills assessment library",
             "Pre-built tests: technical, cognitive, language proficiency, personality",
             "Agencies must validate candidate claims before presenting to clients"],
            ["Coding assessments",
             "Browser-based IDE with 30+ languages, test case execution, anti-cheating",
             "Table stakes for technical recruiting agencies"],
            ["Structured scorecards",
             "Customisable rating rubrics aligned to job requirements with weighted scoring",
             "Eliminates gut-feel evaluations; improves consistency"],
            ["Anti-cheating & proctoring",
             "Tab-switch detection, copy-paste monitoring, AI-response detection, optional webcam proctoring",
             "Cheating rates have doubled since 2024"],
        ],
        col_widths_cm=[3.8, 6.2, 6.0],
    )

    add_heading(doc, "4.6 Reporting, analytics & compliance", level=2)
    add_feature_table(
        doc,
        headers=["Feature", "Description", "Why it matters"],
        rows=[
            ["Pipeline analytics",
             "Real-time metrics: time-to-fill, time-in-stage, source effectiveness, pass-through",
             "Agencies sell speed; they need data to retain clients"],
            ["Client reporting",
             "Exportable, branded reports: pipeline status, candidate summaries, activity logs",
             "Weekly client updates are standard; manual assembly is painful"],
            ["Compliance & audit trail",
             "Immutable logs, consent management (GDPR, DPDPA), data retention & deletion",
             "Regulatory exposure is growing across US and India"],
            ["DEI & bias reporting",
             "Adverse impact analysis, demographic pass-through rates, bias audit exports",
             "Enterprise clients require agencies to demonstrate fair hiring"],
        ],
        col_widths_cm=[3.8, 6.2, 6.0],
    )

    add_heading(doc, "4.7 Integrations & platform architecture", level=2)
    add_feature_table(
        doc,
        headers=["Feature", "Description", "Why it matters"],
        rows=[
            ["ATS/HRIS connectors",
             "Workday, Greenhouse, Lever, iCIMS, SAP SF, BambooHR, Zoho Recruit, Keka, Darwinbox",
             "Agencies push candidates into client ATS; bidirectional sync is critical"],
            ["Open API",
             "RESTful API with webhooks for custom integrations, data export, automation",
             "Larger agencies build custom tooling around their core platform"],
            ["Communication integrations",
             "Gmail, Outlook, Twilio SMS, WhatsApp Business API, calendar sync",
             "All candidate touchpoints must flow through the platform for tracking"],
            ["SSO & security",
             "SAML/OAuth SSO, 2FA, SOC 2 compliance, encryption at rest and in transit",
             "Enterprise clients require certifications before procurement"],
        ],
        col_widths_cm=[3.8, 6.2, 6.0],
    )

    add_heading(doc, "4.8 AI capabilities \u2014 essential (launch)", level=2)
    add_feature_table(
        doc,
        headers=["AI feature", "How it works", "User value", "Benchmark"],
        rows=[
            ["AI conversational interviewer",
             "LLM-powered agent conducts adaptive two-way interviews (video, voice, chat) with real-time follow-ups",
             "Screens candidates before human review; 24/7 across timezones (pre-screening, not full rounds)",
             "Apriora/Alex, HackerEarth Tara, FloCareer NIVO"],
            ["AI candidate scoring",
             "Evaluates responses, assessments, and CV data against requirements; explainable composite score",
             "Ranked shortlists instead of manual review; visible rationale",
             "HireVue, Sapia.ai, SHL"],
            ["Resume parsing & matching",
             "NLP-based skills, experience, and qualification extraction with relevance scoring",
             "Surfaces existing DB candidates before external sourcing",
             "iMocha, HiredScore/Workday"],
            ["AI-generated questions",
             "Ingests JD and produces structured question sets: technical, behavioural, situational",
             "Eliminates hours of prep; ensures consistency",
             "HireVue, Apriora/Alex"],
            ["AI interview summaries",
             "Post-interview summaries from transcripts: strengths, concerns, skill gaps, next steps",
             "Structured debriefs instead of raw transcripts",
             "BrightHire/Zoom, Metaview"],
            ["Automated screening chatbot",
             "AI chatbot (text, WhatsApp, SMS) screens availability, salary, visa, notice period, qualifications",
             "Handles 100s of inbound applications per role without recruiter time",
             "Paradox Olivia, Humanly"],
            ["AI bias detection",
             "Monitors scoring for adverse impact across protected categories; audit-ready reports",
             "Built-in compliance with NYC LL144 and EU AI Act",
             "HiredScore, BrightHire"],
        ],
        col_widths_cm=[3.5, 5.0, 4.5, 3.0],
    )

    add_heading(doc, "4.9 AI capabilities \u2014 high-value (roadmap)", level=2)
    add_paragraph(
        doc,
        "The source specification references a Phase 2 tier of high-value AI "
        "features but does not enumerate them. The following are proposed "
        "additions and must be confirmed before Phase 2 planning is finalised.",
        italic=True, color=MUTED,
    )
    add_bullets(doc, [
        "Candidate-job fit forecasting using placement outcome data (retention, performance proxies).",
        "AI-assisted client intake that turns unstructured JDs and emails into structured requisitions.",
        "Agentic sourcing: autonomous outreach across LinkedIn / email with human-in-the-loop approval.",
        "Conversational analytics (\u201cask the data\u201d) over pipeline, placement, and revenue metrics.",
    ])


def build_non_functional(doc: Document) -> None:
    add_heading(doc, "5. Non-functional requirements", level=1)
    add_feature_table(
        doc,
        headers=["Category", "Requirement"],
        rows=[
            ["Availability", "99.9% monthly uptime SLA for Phase 2+; 99.5% during Phase 1 beta."],
            ["Performance", "P95 API latency < 400 ms; AI inference results streamed within 3 s."],
            ["Data residency", "Separate US and India data regions; no cross-border PII transfer by default."],
            ["Security", "SOC 2 Type II by end of Phase 3; SSO, MFA, encryption at rest (AES-256) and in transit (TLS 1.2+)."],
            ["Privacy", "GDPR, CCPA, DPDPA compliant; consent management and right-to-erasure workflows."],
            ["Responsible AI", "Model cards, adverse-impact audits per tenant, human-in-the-loop for rejection decisions."],
            ["Accessibility", "WCAG 2.1 AA for recruiter, candidate, and client portals."],
            ["Scalability", "Support 10k concurrent candidate sessions and 1M candidates per tenant."],
            ["Observability", "Centralised logging, tracing, and audit trail with 12-month retention minimum."],
        ],
        col_widths_cm=[3.8, 12.2],
    )


def build_out_of_scope(doc: Document) -> None:
    add_heading(doc, "6. Out of scope (first 18 months)", level=1)
    add_bullets(doc, [
        "Full payroll, background checks, and onboarding \u2014 delivered via integrations, not in-platform.",
        "Job board aggregation / programmatic job advertising.",
        "Native mobile applications (mobile-web responsive only in Phases 1\u20132).",
        "Full HRIS replacement \u2014 the platform is recruiting-centric, not employee-lifecycle.",
        "On-premise deployments.",
    ])


# ---------------------------------------------------------------------------
# Phasing
# ---------------------------------------------------------------------------

PHASE1 = {
    "name": "Phase 1 \u2014 MVP & launch",
    "window": "Months 0\u20136",
    "theme": "Prove agency-first workflow with a credible AI layer",
    "objectives": [
        "Ship a workable end-to-end recruiting workflow for a single-region (US or India) pilot cohort.",
        "Deliver essential AI features that justify the AI-first positioning.",
        "Secure 3\u20135 design-partner agencies with live placements.",
    ],
    "features": [
        ["Multi-client workspace management", "Client workspaces, role-based access"],
        ["Candidate pipeline", "Universal candidate DB, customisable pipelines, rich profiles, communication hub (email)"],
        ["Job management", "Job requisition builder, application parsing"],
        ["Scheduling", "Automated scheduling (Google + Outlook), reminders (email + SMS)"],
        ["Assessment", "Structured scorecards"],
        ["Analytics", "Core pipeline analytics dashboard"],
        ["Integrations", "SSO (Google, Microsoft), Gmail + Outlook email/calendar"],
        ["AI \u2014 essential", "Resume parsing & matching, AI-generated questions, AI interview summaries"],
        ["Security & compliance", "Encryption at rest/in transit, audit log foundation, GDPR/DPDPA consent baseline"],
    ],
    "exit": [
        "3+ design partners actively placing candidates via the platform.",
        "Time-to-shortlist median < 72 hours in pilot accounts.",
        "Zero Sev-1 incidents in the final month of the phase.",
    ],
}

PHASE2 = {
    "name": "Phase 2 \u2014 AI differentiation & assessment depth",
    "window": "Months 6\u201312",
    "theme": "Own the interview + screening loop; prove defensible AI",
    "objectives": [
        "Move from \u201cAI-assisted\u201d to \u201cAI-executed\u201d screening for the majority of inbound candidates.",
        "Make the platform sticky through assessments, video, and multi-channel communication.",
        "Expand commercially to paid GA across both US and India.",
    ],
    "features": [
        ["Multi-client workspace management", "White-label portals, multi-entity billing"],
        ["Candidate pipeline", "Duplicate detection, submission tracking, WhatsApp + SMS in communication hub"],
        ["Scheduling", "Panel coordination, feedback collection"],
        ["Assessment", "Live video interviewing, skills assessment library, coding assessments, anti-cheating (tab-switch + copy-paste + AI-response)"],
        ["Analytics", "Branded client reporting, weekly digest automation"],
        ["Integrations", "Top-3 ATS connectors (Greenhouse, Lever, Workday), Twilio SMS, WhatsApp Business API"],
        ["AI \u2014 essential", "AI conversational interviewer (chat & voice), AI candidate scoring, automated screening chatbot"],
        ["Responsible AI", "AI bias detection v1 (adverse-impact dashboards)"],
        ["Compliance", "SOC 2 Type I, GDPR/DPDPA full audit trail, consent & retention workflows"],
    ],
    "exit": [
        "AI screening coverage > 50% of inbound applications.",
        "Paying customers in both US and India regions.",
        "SOC 2 Type I report issued.",
    ],
}

PHASE3 = {
    "name": "Phase 3 \u2014 Enterprise scale & advanced AI",
    "window": "Months 12\u201318",
    "theme": "Win enterprise procurement and compound AI advantage",
    "objectives": [
        "Close enterprise agency deals that require SOC 2 Type II and deep ATS integrations.",
        "Ship high-value AI that the market has not yet productised.",
        "Harden the platform for multi-region scale and regulatory scrutiny.",
    ],
    "features": [
        ["Integrations", "Remaining ATS/HRIS connectors (iCIMS, SAP SF, BambooHR, Zoho, Keka, Darwinbox); Open API + webhooks"],
        ["Assessment", "Video interview proctoring (webcam), advanced anti-cheating signals"],
        ["Analytics", "DEI & bias reporting, forecasting, conversational analytics"],
        ["AI \u2014 high-value", "Candidate-job fit forecasting, AI-assisted client intake, agentic sourcing (human-in-the-loop)"],
        ["Responsible AI", "Per-tenant bias audit reports; NYC LL144 & EU AI Act artefacts"],
        ["Platform", "Multi-region deployment (US + India), zero-downtime migrations, tenant isolation hardening"],
        ["Compliance", "SOC 2 Type II, ISO 27001 readiness, data residency contracts"],
    ],
    "exit": [
        "First enterprise logos (1k+ seats) live.",
        "SOC 2 Type II achieved.",
        "Gross revenue retention > 95%.",
    ],
}


def build_phase(doc: Document, phase: dict) -> None:
    add_heading(doc, phase["name"], level=2)
    add_paragraph(doc, f"Timeline: {phase['window']}", bold=True)
    add_paragraph(doc, f"Theme: {phase['theme']}", italic=True, color=MUTED)

    add_paragraph(doc, "Objectives", bold=True)
    add_bullets(doc, phase["objectives"])

    add_paragraph(doc, "In-scope features", bold=True)
    add_feature_table(
        doc,
        headers=["Pillar", "Features delivered"],
        rows=phase["features"],
        col_widths_cm=[5.0, 11.0],
    )

    add_paragraph(doc, "Exit criteria", bold=True)
    add_bullets(doc, phase["exit"])
    doc.add_paragraph()


def build_phasing(doc: Document) -> None:
    add_heading(doc, "7. Proposed delivery phases", level=1)
    add_paragraph(
        doc,
        "The roadmap below is a proposal, not a final decision. Phase content "
        "can be moved between phases based on engineering capacity, design-partner "
        "feedback, and commercial priorities. The guiding principle is: Phase 1 "
        "earns the right to exist, Phase 2 earns the right to charge premium "
        "pricing, Phase 3 earns the right to sell to the enterprise.",
        italic=True, color=MUTED,
    )
    for phase in (PHASE1, PHASE2, PHASE3):
        build_phase(doc, phase)


def build_risks(doc: Document) -> None:
    add_heading(doc, "8. Risks & mitigations", level=1)
    add_feature_table(
        doc,
        headers=["Risk", "Impact", "Mitigation"],
        rows=[
            ["LLM hallucinations in candidate scoring",
             "Legal exposure, unfair rejections",
             "Human-in-the-loop for all reject decisions; explainability required for scores; adverse-impact monitoring"],
            ["Cheating on AI-administered assessments",
             "Loss of credibility with clients",
             "Layered anti-cheating (tab-switch, copy-paste, AI-response detection, optional proctoring) with visible trust signals"],
            ["Regulatory divergence between US and India",
             "Compliance rework, blocked deals",
             "Region-specific deployments, per-tenant policy configs, legal review before market entry"],
            ["Integration fragility with legacy ATS",
             "Churn risk at enterprise tier",
             "Prioritise top-3 ATS in Phase 2, invest in sync health monitoring, publish SLAs"],
            ["Model cost escalation",
             "Gross margin compression",
             "Small-model-first architecture; cache prompts; per-tenant usage caps"],
            ["Commoditisation by ATS incumbents",
             "Pricing pressure",
             "Lean into agency-specific workflows, multi-client data model, and placement-centric analytics as moats"],
        ],
        col_widths_cm=[4.0, 4.0, 8.0],
    )


def build_open_questions(doc: Document) -> None:
    add_heading(doc, "9. Open questions & decisions required", level=1)
    add_bullets(doc, [
        "Which region (US or India) is the Phase 1 pilot region? This drives compliance, integrations, and comms channel priorities.",
        "What pricing model should the platform support first: per-seat, per-placement, or per-interview?",
        "Which design partners will co-develop Phase 1? Their workflows will heavily influence the MVP surface area.",
        "Confirm the Phase 2 high-value AI list \u2014 the source spec hints at this tier but does not enumerate it.",
        "What is the tolerance for build-vs-buy on video infrastructure (Phase 2)? This is a major cost and complexity decision.",
    ])


def build_limitations(doc: Document) -> None:
    add_heading(doc, "10. Limitations of this PRD", level=1)
    add_bullets(doc, [
        "The source feature specification ends at section 2.1; sections 2.2 (high-value AI) and 3 (good-to-have features) are referenced but not included. High-value AI proposals in section 4.9 are therefore inferred and must be validated.",
        "Success metric targets are based on industry benchmarks, not internal data; they should be re-baselined after Phase 1 pilots.",
        "Engineering effort estimates are not included \u2014 phases are scoped by capability, not staffing. An engineering capacity plan is a required next artefact.",
        "This PRD does not include detailed wireframes, data model, or API contracts; those are per-feature design artefacts.",
    ])


def build_next_steps(doc: Document) -> None:
    add_heading(doc, "11. Next steps", level=1)
    add_bullets(doc, [
        "Review and confirm personas, success metrics, and phase boundaries with product and engineering leads.",
        "Resolve the open questions in section 9, especially pilot region and design partners.",
        "Produce an engineering capacity and staffing plan mapped to the three phases.",
        "Break Phase 1 into epics and user stories; draft per-feature PRD addenda for the top 3 features.",
        "Schedule a responsible-AI review covering scoring, bias detection, and audit workflows before Phase 1 build begins.",
    ])


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    build_cover(doc)
    build_executive_summary(doc)
    build_problem_and_users(doc)
    build_goals_and_metrics(doc)
    build_scope(doc)
    build_non_functional(doc)
    build_out_of_scope(doc)
    build_phasing(doc)
    build_risks(doc)
    build_open_questions(doc)
    build_limitations(doc)
    build_next_steps(doc)

    project_root = Path(__file__).resolve().parent.parent
    out_path = project_root / "AIRIS_PRD_v1.0.docx"
    doc.save(out_path)
    print(f"PRD written to: {out_path}")


if __name__ == "__main__":
    main()

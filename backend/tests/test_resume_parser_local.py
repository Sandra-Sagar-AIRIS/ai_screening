from __future__ import annotations

from pathlib import Path

from app.services.resume_parser import parse_resume_file


def test_parse_resume_file_docx_extracts_text(tmp_path: Path) -> None:
    from docx import Document

    file_path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Skills")
    doc.add_paragraph("Python, FastAPI, PostgreSQL")
    doc.add_paragraph("Experience")
    doc.add_paragraph("5 years experience")
    doc.save(str(file_path))

    parsed = parse_resume_file(file_path)
    assert "Jane Doe" in parsed.text
    assert parsed.sections.get("skills")
    assert "Python" in parsed.text


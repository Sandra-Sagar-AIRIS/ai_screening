"""Local-first resume text extractor.

Tries pdfplumber for PDFs (better layout), falls back to PyMuPDF, and uses
python-docx for DOCX files. Returns raw text plus a best-effort section
breakdown that downstream extraction can consume without requiring an LLM.
"""
from __future__ import annotations

import logging
import re
import time
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


SUPPORTED_EXTS: frozenset[str] = frozenset({".pdf", ".docx"})

_SECTION_PATTERNS: tuple[tuple[str, re.Pattern[str]], ...] = (
    ("summary", re.compile(r"^\s*(professional\s+summary|summary|profile|objective|about)\s*:?\s*$", re.IGNORECASE)),
    ("experience", re.compile(r"^\s*(work\s+experience|professional\s+experience|experience|employment(?:\s+history)?)\s*:?\s*$", re.IGNORECASE)),
    ("education", re.compile(r"^\s*(education|academic(?:\s+background)?|qualifications)\s*:?\s*$", re.IGNORECASE)),
    ("skills", re.compile(r"^\s*(skills|technical\s+skills|core\s+competenc(?:y|ies)|key\s+skills)\s*:?\s*$", re.IGNORECASE)),
    ("certifications", re.compile(r"^\s*(certifications?|certificates?|licenses?)\s*:?\s*$", re.IGNORECASE)),
    ("projects", re.compile(r"^\s*(projects?|key\s+projects?)\s*:?\s*$", re.IGNORECASE)),
)

_SECTION_KEYS: tuple[str, ...] = tuple(k for k, _ in _SECTION_PATTERNS)


@dataclass(slots=True)
class ParsedResume:
    """Result of local resume parsing.

    `text` is the full extracted text. `sections` maps a canonical section name
    (e.g. "experience", "education") to the lines that fall under that header in
    the document. `parser` tells you which backend produced the text so callers
    can apply per-parser confidence in scoring.
    """

    text: str
    sections: dict[str, str] = field(default_factory=dict)
    parser: str = "unknown"
    page_count: int = 0


class ResumeParseError(Exception):
    """Raised when none of the local extractors can produce usable text."""


def parse_resume_file(file_path: Path | str) -> ParsedResume:
    """Extract text and rough sections from a local resume file."""
    t0 = time.monotonic()
    path = Path(file_path)
    if not path.exists():
        raise ResumeParseError(f"resume file not found: {path}")
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text, parser, page_count = _extract_pdf(path)
    elif suffix == ".docx":
        text, parser, page_count = _extract_docx(path)
    else:
        # Best-effort plain text fallback so the pipeline never crashes on
        # an unexpected file type — but the caller's MIME guard should
        # already have rejected anything but PDF/DOCX upstream.
        text, parser, page_count = _extract_plaintext(path), "plaintext", 0

    text = _normalize_text(text)
    if not text.strip():
        raise ResumeParseError(f"no extractable text in resume: {path.name}")
    sections = _split_sections(text)
    logger.info(
        "ats.resume.parse.completed",
        extra={
            "ats_phase": "resume_parse",
            "file_name": path.name,
            "parser": parser,
            "page_count": page_count,
            "duration_ms": int((time.monotonic() - t0) * 1000),
            "text_chars": len(text),
        },
    )
    return ParsedResume(text=text, sections=sections, parser=parser, page_count=page_count)


def _extract_pdf(path: Path) -> tuple[str, str, int]:
    # pdfplumber preserves spacing/line layout much better than pypdf, which
    # matters for the regex section detector and skill matching downstream.
    try:
        t_pdfplumber = time.monotonic()
        import pdfplumber  # type: ignore

        chunks: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text:
                    chunks.append(page_text)
            page_count = len(pdf.pages)
        text = "\n".join(chunks)
        if text.strip():
            logger.info(
                "ats.pdf.extract.completed",
                extra={
                    "ats_phase": "pdf_extract",
                    "file_name": path.name,
                    "extractor": "pdfplumber",
                    "duration_ms": int((time.monotonic() - t_pdfplumber) * 1000),
                    "page_count": page_count,
                },
            )
            return text, "pdfplumber", page_count
    except Exception as exc:  # noqa: BLE001
        logger.warning("pdfplumber failed for %s: %s", path.name, exc)

    try:
        t_pymupdf = time.monotonic()
        import fitz  # type: ignore  # PyMuPDF

        with fitz.open(str(path)) as doc:
            chunks = [page.get_text("text") or "" for page in doc]
            page_count = doc.page_count
        text = "\n".join(chunks)
        if text.strip():
            logger.info(
                "ats.pdf.extract.completed",
                extra={
                    "ats_phase": "pdf_extract",
                    "file_name": path.name,
                    "extractor": "pymupdf",
                    "duration_ms": int((time.monotonic() - t_pymupdf) * 1000),
                    "page_count": page_count,
                },
            )
            return text, "pymupdf", page_count
    except Exception as exc:  # noqa: BLE001
        logger.warning("pymupdf failed for %s: %s", path.name, exc)

    # Last-resort fallback. pypdf is already a dependency and is sometimes the
    # only thing that opens a damaged PDF.
    try:
        t_pypdf = time.monotonic()
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        chunks = [(page.extract_text() or "") for page in reader.pages]
        text = "\n".join(chunks)
        pages = len(reader.pages)
        logger.info(
            "ats.pdf.extract.completed",
            extra={
                "ats_phase": "pdf_extract",
                "file_name": path.name,
                "extractor": "pypdf",
                "duration_ms": int((time.monotonic() - t_pypdf) * 1000),
                "page_count": pages,
            },
        )
        return text, "pypdf", pages
    except Exception as exc:  # noqa: BLE001
        logger.warning("pypdf failed for %s: %s", path.name, exc)
        return "", "none", 0


def _extract_docx(path: Path) -> tuple[str, str, int]:
    try:
        from docx import Document  # type: ignore

        doc = Document(str(path))
        # Paragraphs cover most resumes; tables are added because some
        # candidates ship resumes with skill matrices in tables.
        chunks: list[str] = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        chunks.append(cell.text)
        return "\n".join(chunks), "python-docx", 0
    except Exception as exc:  # noqa: BLE001
        logger.warning("python-docx failed for %s: %s", path.name, exc)
        return "", "none", 0


def _extract_plaintext(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:  # noqa: BLE001
        logger.warning("plaintext read failed for %s: %s", path.name, exc)
        return ""


def _normalize_text(text: str) -> str:
    # Drop NUL bytes (Postgres rejects them in jsonb/text), collapse weird
    # whitespace, and strip page-break artifacts.
    text = text.replace("\x00", " ").replace("\u200b", " ").replace("\f", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _split_sections(text: str) -> dict[str, str]:
    """Best-effort section split using header regexes.

    We never claim 100% accuracy here — we just need a hint that "this block of
    lines is the experience section" so the extractor can apply per-section
    rules (e.g. job titles only inside experience).
    """
    sections: dict[str, list[str]] = {key: [] for key in _SECTION_KEYS}
    current: str | None = None
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            if current is not None:
                sections[current].append("")
            continue
        matched_section: str | None = None
        for key, pattern in _SECTION_PATTERNS:
            if pattern.match(line):
                matched_section = key
                break
        if matched_section is not None:
            current = matched_section
            continue
        if current is not None:
            sections[current].append(line)
    return {key: "\n".join(lines).strip() for key, lines in sections.items() if lines}
